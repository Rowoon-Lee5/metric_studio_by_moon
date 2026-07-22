"""Build a publication-ready Word packet for Metric Studio chapter 2, parts 2-4."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DRAFTS = ROOT / "blog_drafts"
ASSETS = ROOT / "blog_assets"
OUTPUT = ROOT / "deliverables" / "문병로_교수의_메트릭_스튜디오_02_2부부터_4부_게시용.docx"

IMAGE_MAP = {
    "02_거래회전율은_관심이_아니었다.md": [
        ("06_book_excerpt_popularity.png", "책 2장 발췌 - 관심과 인기주의"),
        ("03_observed_news_attention.png", "그림 1. 원천 뉴스량별 모멘텀 Rank IC"),
    ],
    "03_모델합의는_알파가_아니었다.md": [
        ("07_book_excerpt_strategy.png", "책 2장 발췌 - 패턴과 운용 전략"),
        ("04_model_consensus_cumulative.png", "그림 2. 모델 합의 포트폴리오의 누적 성과"),
    ],
    "04_시장은_신호의_실패방식으로_드러난다.md": [
        ("08_book_excerpt_mixed_results.png", "책 2장 발췌 - 상반된 시장 관찰"),
        ("05_model_failure_coherence.png", "그림 3. 동시에 실패한 모델 수와 다음 달 시장 수익률"),
    ],
}


def set_font(run, size=11, bold=None, color=None, italic=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(doc, with_footer=True):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Malgun Gothic"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(89, 89, 89)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.line_spacing = 1.15

    if with_footer:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("문병로 교수의 메트릭 스튜디오 02 | 연구 노트")
        set_font(run, 8.5, color=(100, 100, 100))


def add_cover(doc, subtitle_text="2장 시장관찰 | ②~④\n관심, 합의, 실패 공분산", description="책의 문장을 요약하지 않고, 실제 한국 주식 데이터로\n가설의 측정 방식과 실패 조건을 검증한 연구 기록"):
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("RESEARCH NOTE SERIES")
    set_font(r, 10, bold=True, color=(122, 90, 0))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("문병로 교수의 메트릭 스튜디오 02")
    set_font(r, 25, bold=True, color=(11, 37, 69))
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(26)
    r = subtitle_para.add_run(subtitle_text)
    set_font(r, 15, color=(43, 81, 99))
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(description)
    set_font(r, 11, italic=True, color=(80, 80, 80))
    doc.add_page_break()


def add_image(doc, filename, caption):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.1))
    p.paragraph_format.keep_with_next = True
    cap = doc.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_font(run, 9, color=(89, 89, 89))


def add_quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, 110, 180, 110, 180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, 10.5, bold=True, color=(31, 58, 95))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_paragraph(doc, text):
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, 11, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, 11)


def add_table(doc, rows):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if len(rows[0]) == 3:
        widths = [Inches(3.4), Inches(1.55), Inches(1.2)]
    elif len(rows[0]) == 4:
        widths = [Inches(1.25), Inches(0.7), Inches(2.0), Inches(2.15)]
    else:
        widths = [Inches(6.15 / len(rows[0]))] * len(rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = widths[idx]
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cells[idx], "F2F4F7")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, 10, bold=(row_idx == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_sources(doc, lines, heading):
    doc.add_heading(heading, level=2)
    for line in lines:
        match = re.match(r"- \[(.*?)\]\((.*?)\)", line)
        if match:
            label, url = match.groups()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{label}: {url}")
            set_font(r, 9.5, color=(55, 55, 55))


def add_article(doc, draft_name):
    text = (DRAFTS / draft_name).read_text(encoding="utf-8").replace("\r\n", "\n")
    images = iter(IMAGE_MAP[draft_name])
    chunks = text.split("\n")
    i = 0
    while i < len(chunks):
        line = chunks[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("**[책 발췌 이미지") or line.startswith("**[이미지"):
            filename, caption = next(images)
            add_image(doc, filename, caption)
        elif line == "---":
            pass
        elif line.startswith("> **"):
            quote_lines = []
            while i < len(chunks) and chunks[i].strip().startswith(">"):
                quote_lines.append(chunks[i].strip().lstrip("> "))
                i += 1
            add_quote(doc, "\n".join(quote_lines).replace("**", ""))
            continue
        elif line.startswith("|"):
            table_lines = []
            while i < len(chunks) and chunks[i].strip().startswith("|"):
                table_lines.append(chunks[i].strip())
                i += 1
            parsed = []
            for tbl_line in table_lines:
                values = [x.strip() for x in tbl_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-+:?", v) for v in values):
                    continue
                parsed.append(values)
            if parsed:
                add_table(doc, parsed)
            continue
        elif line in {"### 재현 자료", "### GitHub에서 확인할 수 있는 것"}:
            source_lines = []
            disclaimer = None
            heading = line[4:]
            i += 1
            while i < len(chunks):
                candidate = chunks[i].strip()
                if candidate.startswith("- ["):
                    source_lines.append(candidate)
                elif candidate.startswith("이 글은 투자 권유"):
                    disclaimer = candidate
                    break
                i += 1
            add_sources(doc, source_lines, heading)
            if disclaimer:
                add_markdown_paragraph(doc, disclaimer)
            continue
        else:
            add_markdown_paragraph(doc, line)
        i += 1


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    for number, draft in enumerate(IMAGE_MAP):
        if number:
            doc.add_page_break()
        add_article(doc, draft)
    doc.core_properties.title = "문병로 교수의 메트릭 스튜디오 02 | 2장 시장관찰 ②~④"
    doc.core_properties.author = "beneficial5"
    doc.core_properties.subject = "한국 주식 데이터 기반 연구 노트"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
