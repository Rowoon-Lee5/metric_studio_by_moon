"""Turn the complete research-package inventory into a reader-facing Word guide."""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_parts_02_04_docx as packet


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "최종_전체_연구_패키지"
SOURCE = PACKAGE / "전체_파일_설명서.md"
OUTPUT = PACKAGE / "연구_자료_전체_설명서.docx"


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("문병로 교수의 메트릭 스튜디오 02")
    packet.set_font(r, 22, bold=True, color=(11, 37, 69))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("연구 자료 전체 설명서 | 코드와 CSV를 열지 않고 읽는 버전")
    packet.set_font(r, 12, color=(80, 80, 80))
    packet.add_quote(doc, "이 문서는 복사·붙여넣기용 블로그 원고가 아니라, 1~4편의 수치와 결론이 어디에서 왔는지 확인하는 읽기용 안내서다. 실제 게시에는 00_게시_자료 안의 Word 원고 네 개만 사용하면 된다.")


def render_markdown(doc, text):
    lines = text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            # The source title is represented by the document title block.
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("|"):
            raw = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw.append(lines[i].strip())
                i += 1
            rows = []
            for candidate in raw:
                cells = [x.strip().strip("`") for x in candidate.strip("|").split("|")]
                if all(re.fullmatch(r":?-+:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            if rows:
                packet.add_table(doc, rows)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[2:])
            packet.set_font(r, 10.5)
        else:
            packet.add_markdown_paragraph(doc, line.replace("`", ""))
        i += 1


def main():
    doc = Document()
    packet.style_document(doc, with_footer=False)
    add_title(doc)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "문병로 교수의 메트릭 스튜디오 02 | 연구 자료 전체 설명서"
    doc.core_properties.author = "beneficial5"
    doc.core_properties.subject = "연구 코드와 검증 결과의 읽기용 안내서"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
