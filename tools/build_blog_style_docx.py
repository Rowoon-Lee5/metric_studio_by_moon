"""Create one Word file per blog post, following the author's existing Naver-post flow."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_parts_02_04_docx as packet


OUTPUTS = {
    "00_연재_소개.md": "문병로_교수의_메트릭_스튜디오_02_00_블로그_게시용.docx",
    "01_알파는_봉우리인가_대륙인가.md": "문병로_교수의_메트릭_스튜디오_02_01_블로그_게시용.docx",
    "02_거래회전율은_관심이_아니었다.md": "문병로_교수의_메트릭_스튜디오_02_02_블로그_게시용.docx",
    "03_모델합의는_알파가_아니었다.md": "문병로_교수의_메트릭_스튜디오_02_03_블로그_게시용.docx",
    "04_시장은_신호의_실패방식으로_드러난다.md": "문병로_교수의_메트릭_스튜디오_02_04_블로그_게시용.docx",
}

packet.IMAGE_MAP = {
    "00_연재_소개.md": [
        ("09_book_excerpt_liquidity.png", "책 2장 발췌 - 유동성 유니버스와 슬리피지"),
    ],
    "01_알파는_봉우리인가_대륙인가.md": [
        ("09_book_excerpt_liquidity.png", "책 2장 발췌 - 유동성 유니버스와 슬리피지"),
        ("01_robustness_map.png", "그림 1. 유동성·비용 조건별로 살아남은 전략 조합"),
        ("02_evidence_decision_boundary.png", "그림 2. 다중검정 보정 뒤의 의사결정 경계"),
        ("10_suspension_screened_smallcap.png", "그림 3. 거래정지 편입 제외 뒤 소형주 전략 성과"),
    ],
    "02_거래회전율은_관심이_아니었다.md": [
        ("06_book_excerpt_popularity.png", "책 2장 발췌 - 관심과 인기주의"),
        ("03_observed_news_attention.png", "그림 1. 원천 뉴스량별 모멘텀 Rank IC"),
    ],
    "03_모델합의는_알파가_아니었다.md": [
        ("07_book_excerpt_strategy.png", "책 2장 발췌 - 패턴과 운용 전략"),
        ("04_model_consensus_cumulative.png", "그림 1. 모델 합의 포트폴리오의 누적 성과"),
    ],
    "04_시장은_신호의_실패방식으로_드러난다.md": [
        ("08_book_excerpt_mixed_results.png", "책 2장 발췌 - 상반된 시장 관찰"),
        ("05_model_failure_coherence.png", "그림 1. 동시에 실패한 모델 수와 다음 달 시장 수익률"),
    ],
}


def make_post(draft_name: str, output_name: str):
    doc = Document()
    packet.style_document(doc, with_footer=False)
    # The existing post begins with its title and personal prose; no report cover,
    # running furniture, or artificial "research packet" preface is inserted.
    doc.styles["Heading 1"].font.size = Pt(17)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(0)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(18)
    packet.add_article(doc, draft_name)
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            break
    output = packet.ROOT / "deliverables" / output_name
    doc.core_properties.title = draft_name.removesuffix(".md")
    doc.core_properties.author = "beneficial5"
    doc.core_properties.subject = "문병로 교수의 메트릭 스튜디오 02 블로그 원고"
    doc.save(output)
    print(output)


def main():
    for draft, output in OUTPUTS.items():
        make_post(draft, output)
    combined = Document()
    packet.style_document(combined, with_footer=False)
    for index, draft in enumerate(OUTPUTS):
        if index:
            combined.add_page_break()
        packet.add_article(combined, draft)
    combined.core_properties.title = "문병로 교수의 메트릭 스튜디오 02 | 블로그 게시용 0~4"
    combined.core_properties.author = "beneficial5"
    combined.core_properties.subject = "문병로 교수의 메트릭 스튜디오 02 블로그 원고"
    combined_output = packet.ROOT / "deliverables" / "문병로_교수의_메트릭_스튜디오_02_00부터_04_블로그_게시용.docx"
    combined.save(combined_output)
    print(combined_output)


if __name__ == "__main__":
    main()
